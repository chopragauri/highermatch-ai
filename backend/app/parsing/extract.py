import io

import pdfplumber
from docx import Document


def clean_ligatures(text: str) -> str:
    import re
    replacements = {
        "(cid:27)": "ff",
        "(cid:28)": "fi",
        "(cid:29)": "fl",
        "(cid:30)": "ffi",
        "(cid:31)": "ffl",
    }
    for cid, rep in replacements.items():
        text = text.replace(cid, rep)
    # Strip any remaining unmapped cids
    text = re.sub(r"\(cid:\d+\)", "", text)
    return text


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
    except Exception:
        pass

    text = "\n".join(text_parts).strip()
    if text:
        return clean_ligatures(text)

    # Fallback: some PDFs pdfplumber can't parse cleanly still work with PyPDF2.
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = [page.extract_text() or "" for page in reader.pages]
    except Exception:
        pass

    return clean_ligatures("\n".join(text_parts).strip())


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    # Many resumes put skills/experience in tables rather than paragraphs.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts).strip()


def extract_text(file_bytes: bytes, mime: str) -> str:
    mime = (mime or "").lower()
    if "pdf" in mime:
        return extract_text_from_pdf(file_bytes)
    if "word" in mime or "docx" in mime or "officedocument" in mime:
        return extract_text_from_docx(file_bytes)
    if "text/plain" in mime:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    # Unknown mime: best-effort try both extractors.
    try:
        text = extract_text_from_pdf(file_bytes)
        if text:
            return text
    except Exception:
        pass
    return extract_text_from_docx(file_bytes)
